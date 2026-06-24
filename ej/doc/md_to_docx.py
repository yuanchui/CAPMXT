#!/usr/bin/env python3
"""Convert project markdown summary to Word (.docx)."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def set_cell_shading(cell, fill: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def add_formatted_runs(paragraph, text: str, base_bold: bool = False) -> None:
    pattern = re.compile(r"\*\*(.+?)\*\*|`([^`]+)`")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            run.bold = base_bold
        if match.group(1):
            run = paragraph.add_run(match.group(1))
            run.bold = True
        else:
            run = paragraph.add_run(match.group(2))
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(9)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.bold = base_bold


def parse_table_row(line: str) -> list[str]:
    line = line.strip().strip("|")
    return [cell.strip() for cell in line.split("|")]


def is_table_sep(line: str) -> bool:
    return bool(re.match(r"^\|\s*[-:]+(\s*\|\s*[-:]+)+\s*\|?\s*$", line.strip()))


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "微软雅黑"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)

    for level, size in [(1, 18), (2, 14), (3, 12)]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    table.autofit = True

    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            cell = table.rows[r_idx].cells[c_idx]
            value = row[c_idx] if c_idx < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            add_formatted_runs(p, value, base_bold=(r_idx == 0))
            for run in p.runs:
                run.font.size = Pt(9)
            if r_idx == 0:
                set_cell_shading(cell, "D9E2F3")


def add_image(doc: Document, md_path: Path, image_ref: str, alt: str = "") -> None:
    img_path = (md_path.parent / image_ref).resolve()
    if not img_path.is_file():
        p = doc.add_paragraph()
        p.add_run(f"[图片缺失: {image_ref}]").italic = True
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=Cm(15.5))
    if alt:
        cap = doc.add_paragraph(alt)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def convert_markdown(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)
    configure_styles(doc)

    i = 0
    in_code = False
    code_lines: list[str] = []
    table_rows: list[list[str]] | None = None
    list_buffer: list[tuple[str, str]] | None = None

    def flush_table() -> None:
        nonlocal table_rows
        if table_rows is not None:
            add_table(doc, table_rows)
            table_rows = None

    def flush_list() -> None:
        nonlocal list_buffer
        if not list_buffer:
            return
        for kind, text in list_buffer:
            p = doc.add_paragraph(style="List Bullet" if kind == "ul" else "List Number")
            add_formatted_runs(p, text)
        list_buffer = None

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
                run.font.size = Pt(9)
                p.paragraph_format.left_indent = Cm(0.5)
                code_lines = []
                in_code = False
            else:
                flush_list()
                flush_table()
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if stripped.startswith("|"):
            if is_table_sep(stripped):
                i += 1
                continue
            if table_rows is None:
                flush_list()
                table_rows = []
            table_rows.append(parse_table_row(stripped))
            i += 1
            continue

        if table_rows is not None:
            flush_table()

        m_img = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)$", stripped)
        if m_img:
            flush_list()
            add_image(doc, md_path, m_img.group(2), m_img.group(1))
            i += 1
            continue

        if stripped in ("---", "***", "___"):
            flush_list()
            doc.add_paragraph()
            i += 1
            continue

        if stripped.startswith("# "):
            flush_list()
            doc.add_heading(stripped[2:].strip(), level=1)
            i += 1
            continue
        if stripped.startswith("## "):
            flush_list()
            doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith("### "):
            flush_list()
            doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue

        if stripped.startswith("> "):
            flush_list()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            add_formatted_runs(p, stripped[2:])
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        m_ul = re.match(r"^[-*]\s+\[([ xX])\]\s+(.+)$", stripped)
        if m_ul:
            if list_buffer is None:
                list_buffer = []
            checked = m_ul.group(1).lower() == "x"
            prefix = "☑ " if checked else "☐ "
            list_buffer.append(("ul", prefix + m_ul.group(2)))
            i += 1
            continue

        m_ul2 = re.match(r"^[-*]\s+(.+)$", stripped)
        if m_ul2:
            if list_buffer is None:
                list_buffer = []
            list_buffer.append(("ul", m_ul2.group(1)))
            i += 1
            continue

        m_ol = re.match(r"^\d+\.\s+(.+)$", stripped)
        if m_ol:
            if list_buffer is None:
                list_buffer = []
            list_buffer.append(("ol", m_ol.group(1)))
            i += 1
            continue

        if stripped == "":
            flush_list()
            i += 1
            continue

        flush_list()
        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip("*"))
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        else:
            p = doc.add_paragraph()
            add_formatted_runs(p, stripped)
        i += 1

    flush_table()
    flush_list()

    doc.save(docx_path)
    print(f"Wrote: {docx_path}")


def main() -> int:
    if len(sys.argv) < 3:
        print("Usage: md_to_docx.py <input.md> <output.docx>")
        return 1
    convert_markdown(Path(sys.argv[1]), Path(sys.argv[2]))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
